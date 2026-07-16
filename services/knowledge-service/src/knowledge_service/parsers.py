"""Document text extraction by format (RAG phase 5a).

Real telecom documentation is PDF (contracts, T&Cs, regulatory texts, SLAs) and DOCX (internal
procedures), not Markdown. Requiring a human to convert every file first is what keeps a corpus
at 5 documents forever.

Each extractor returns plain text with **paragraph breaks preserved as blank lines**, because
that is the boundary the chunker splits on: emit one run-on blob and every chunk cuts mid-step.
Structured formats (CSV/JSON) are rendered as one readable paragraph per record so a row stays
retrievable as a unit instead of being sliced across chunks.

Extraction never guesses: an unreadable file raises, so ingestion records a failed job rather
than indexing an empty document that would answer questions with silence.
"""
from __future__ import annotations

import csv
import io
import json
import logging
import unicodedata

logger = logging.getLogger(__name__)

# Formats the pipeline can turn into text. minio_store filters the bucket on these.
SUPPORTED_SUFFIXES = (".md", ".markdown", ".txt", ".pdf", ".docx", ".csv", ".json")


class ParseError(RuntimeError):
    """A document could not be turned into text. Never degrade to an empty document."""


def normalize_unicode(text: str) -> str:
    """NFC-normalize and canonicalize newlines.

    NFC matters for Arabic and accented French: identical-looking content in different Unicode
    compositions would otherwise tokenize differently (worse retrieval).
    """
    text = unicodedata.normalize("NFC", text)
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def _decode(raw: bytes) -> str:
    return raw.decode("utf-8", errors="replace")


def _extract_text(raw: bytes) -> str:
    return _decode(raw)


def _extract_pdf(raw: bytes) -> str:
    """Text per page, pages separated by blank lines."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ParseError(f"pypdf is not installed: {exc}") from exc
    try:
        reader = PdfReader(io.BytesIO(raw))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as exc:
        raise ParseError(f"unreadable PDF: {exc}") from exc
    body = "\n\n".join(page for page in pages if page)
    if not body.strip():
        # A scanned PDF has pages but no text layer. Indexing it would create a document that
        # matches nothing; say so instead of pretending it ingested.
        raise ParseError("PDF contains no extractable text (scanned image? OCR is required)")
    return body


def _extract_docx(raw: bytes) -> str:
    """Paragraphs and table cells, in document order."""
    try:
        import docx
    except ImportError as exc:
        raise ParseError(f"python-docx is not installed: {exc}") from exc
    try:
        document = docx.Document(io.BytesIO(raw))
    except Exception as exc:
        raise ParseError(f"unreadable DOCX: {exc}") from exc

    blocks: list[str] = [
        para.text.strip() for para in document.paragraphs if para.text and para.text.strip()
    ]
    # Telecom procedures put tariffs and step tables in tables; dropping them loses the answer.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    body = "\n\n".join(blocks)
    if not body.strip():
        raise ParseError("DOCX contains no text")
    return body


def _extract_csv(raw: bytes) -> str:
    """One readable paragraph per row: `column: value; column: value`.

    Rendering rows as prose (rather than raw CSV) gives the embedder real words to match a
    caller's question against, and keeps each record whole through chunking.
    """
    text = _decode(raw)
    try:
        sample = text[:4096]
        dialect = csv.Sniffer().sniff(sample) if sample.strip() else csv.excel
    except csv.Error:
        dialect = csv.excel
    try:
        reader = csv.DictReader(io.StringIO(text), dialect=dialect)
        rows = list(reader)
    except Exception as exc:
        raise ParseError(f"unreadable CSV: {exc}") from exc
    if not rows:
        raise ParseError("CSV has no data rows")

    blocks: list[str] = []
    for row in rows:
        parts = [
            f"{(key or '').strip()}: {str(value).strip()}"
            for key, value in row.items()
            if key and value is not None and str(value).strip()
        ]
        if parts:
            blocks.append("; ".join(parts) + ".")
    if not blocks:
        raise ParseError("CSV produced no readable rows")
    return "\n\n".join(blocks)


def _flatten_json(value, prefix: str = "") -> list[str]:
    """Depth-first `a.b.c: value` lines; lists are indexed."""
    lines: list[str] = []
    if isinstance(value, dict):
        for key, item in value.items():
            lines.extend(_flatten_json(item, f"{prefix}.{key}" if prefix else str(key)))
    elif isinstance(value, list):
        for index, item in enumerate(value):
            lines.extend(_flatten_json(item, f"{prefix}[{index}]"))
    else:
        if value is not None and str(value).strip():
            lines.append(f"{prefix}: {value}")
    return lines


def _extract_json(raw: bytes) -> str:
    """Flatten to `path: value` lines; a top-level array becomes one paragraph per element."""
    try:
        data = json.loads(_decode(raw))
    except Exception as exc:
        raise ParseError(f"invalid JSON: {exc}") from exc
    if isinstance(data, list):
        blocks = ["\n".join(_flatten_json(item)) for item in data]
        body = "\n\n".join(block for block in blocks if block.strip())
    else:
        body = "\n".join(_flatten_json(data))
    if not body.strip():
        raise ParseError("JSON produced no readable content")
    return body


_EXTRACTORS = {
    ".md": _extract_text,
    ".markdown": _extract_text,
    ".txt": _extract_text,
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".csv": _extract_csv,
    ".json": _extract_json,
}


def suffix_of(key: str) -> str:
    """Lowercase file extension of an object key ('' when it has none)."""
    name = key.rsplit("/", 1)[-1]
    return f".{name.rsplit('.', 1)[-1].lower()}" if "." in name else ""


def is_supported(key: str) -> bool:
    return suffix_of(key) in _EXTRACTORS


def extract_text(key: str, raw: bytes) -> str:
    """Turn stored bytes into normalized plain text. Raises ParseError on anything unusable."""
    suffix = suffix_of(key)
    extractor = _EXTRACTORS.get(suffix)
    if extractor is None:
        raise ParseError(f"unsupported format {suffix!r} for {key!r}")
    if not raw:
        raise ParseError(f"{key!r} is empty")
    return normalize_unicode(extractor(raw))
